#!/usr/bin/env python3
"""Generate Distribution Prime Organization SOP as a branded Word document."""

from __future__ import annotations

import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docxcompose.composer import Composer

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
LOGO = DOCS / "assets" / "distribution-prime-logo.png"
SOURCE_MD = DOCS / "Distribution-Prime-Organization-SOP.md"
OUTPUT = DOCS / "Distribution-Prime-Organization-SOP.docx"


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)


def build_cover() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo.add_run()
        run.add_picture(str(LOGO), width=Inches(1.4))

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("Distribution Prime")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run("Organization Standard Operating Procedure")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    meta = [
        ("Document type", "Organization operating procedure"),
        ("Document version", "3.0"),
        ("Effective date", "March 2025"),
        ("Application", "Distribution Prime"),
        ("Application URL", "https://distribution-prime.pages.dev"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(meta):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for cell in table.rows[i].cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()
    p_org = doc.add_paragraph("Organization record (complete before distribution)")
    p_org.runs[0].bold = True
    p_org.runs[0].font.size = Pt(11)

    org_fields = [
        "Organization name",
        "Workspace ID",
        "Document owner",
        "Approved by",
        "Next review date",
    ]
    org_table = doc.add_table(rows=len(org_fields), cols=2)
    org_table.style = "Table Grid"
    for i, label in enumerate(org_fields):
        org_table.rows[i].cells[0].text = label
        org_table.rows[i].cells[1].text = "_________________________________"

    doc.add_page_break()
    p_toc = doc.add_paragraph("Table of Contents")
    p_toc.runs[0].bold = True
    p_toc.runs[0].font.size = Pt(14)
    add_toc(doc.add_paragraph())
    doc.add_page_break()
    return doc


def pandoc_body(src: Path, dest: Path) -> None:
    cmd = [
        "pandoc",
        str(src),
        "-o",
        str(dest),
        "--from",
        "markdown",
        "--to",
        "docx",
        "--toc",
        "--toc-depth=3",
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(result.stderr or result.stdout or "pandoc failed")


def main() -> int:
    if not SOURCE_MD.exists():
        print(f"Missing source: {SOURCE_MD}", file=sys.stderr)
        return 1
    if not LOGO.exists():
        print(f"Warning: logo not found at {LOGO}", file=sys.stderr)

    with tempfile.TemporaryDirectory() as tmp:
        body_path = Path(tmp) / "body.docx"
        pandoc_body(SOURCE_MD, body_path)

        cover = build_cover()
        composer = Composer(cover)
        composer.append(Document(str(body_path)))
        composer.save(str(OUTPUT))

    print(f"Created: {OUTPUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
