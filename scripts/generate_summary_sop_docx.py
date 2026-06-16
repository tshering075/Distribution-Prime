"""Generate Distribution Prime Summary SOP (visual + logo) as Microsoft Word (.docx)."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Distribution-Prime-Summary-SOP-v1.2.docx"
ASSETS = ROOT / "docs" / "assets" / "sop-visuals"
LOGO = ROOT / "docs" / "assets" / "distribution-prime-logo.png"
LOGO_FALLBACK = ROOT / "public" / "distribution-prime-icon-512.png"

# Brand colors
BLUE = "#0d47a1"
BLUE_LIGHT = "#1565c0"
RED = "#e40521"
GREEN = "#2e7d32"
AMBER = "#f9a825"
GREY = "#eceff1"
WHITE = "#ffffff"


def resolve_logo() -> Path:
    if LOGO.exists():
        return LOGO
    return LOGO_FALLBACK


def try_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def draw_rounded_rect(draw: ImageDraw.ImageDraw, xy, fill: str, outline: str | None = None, radius: int = 12) -> None:
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=2)


def save_diagram(name: str, image: Image.Image) -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    path = ASSETS / name
    image.save(path, "PNG")
    return path


def diagram_architecture() -> Path:
    w, h = 900, 380
    img = Image.new("RGB", (w, h), WHITE)
    draw = ImageDraw.Draw(img)
    title_f = try_font(22, True)
    box_f = try_font(15, True)
    small_f = try_font(13)

    draw.text((w // 2, 28), "Multi-tenant data isolation", fill=BLUE, font=title_f, anchor="mm")

    # Platform layer
    draw_rounded_rect(draw, (40, 60, w - 40, 110), fill=BLUE, radius=14)
    draw.text((w // 2, 85), "Distribution Prime (shared application)", fill=WHITE, font=box_f, anchor="mm")

    # Supabase
    draw_rounded_rect(draw, (120, 140, w - 120, 200), fill=GREY, outline=BLUE)
    draw.text((w // 2, 170), "Supabase database — every row tagged with organization_id + RLS", fill="#333", font=small_f, anchor="mm")

    # Three org boxes
    box_w = 220
    gap = 40
    start_x = (w - (3 * box_w + 2 * gap)) // 2
    labels = [
        ("Workspace A", "Orders · Distributors\nTargets · Stock"),
        ("Workspace B", "Orders · Distributors\nTargets · Stock"),
        ("Workspace C", "Orders · Distributors\nTargets · Stock"),
    ]
    y0, y1 = 230, 350
    for i, (title, body) in enumerate(labels):
        x0 = start_x + i * (box_w + gap)
        x1 = x0 + box_w
        fill = "#e3f2fd" if i == 0 else "#f3e5f5" if i == 1 else "#e8f5e9"
        draw_rounded_rect(draw, (x0, y0, x1, y1), fill=fill, outline=BLUE_LIGHT)
        draw.text((x0 + box_w // 2, y0 + 28), title, fill=BLUE, font=box_f, anchor="mm")
        draw.multiline_text((x0 + box_w // 2, y0 + 70), body, fill="#444", font=small_f, anchor="mm", align="center")

    return save_diagram("architecture.png", img)


def diagram_order_flow() -> Path:
    w, h = 920, 260
    img = Image.new("RGB", (w, h), WHITE)
    draw = ImageDraw.Draw(img)
    title_f = try_font(22, True)
    box_f = try_font(13, True)

    draw.text((w // 2, 24), "Order lifecycle", fill=BLUE, font=title_f, anchor="mm")

    steps = [
        ("1. Pending", "Distributor\nplaces order", AMBER),
        ("2. Sent", "Admin emails\nfor approval", "#90caf9"),
        ("3. Approved", "Admin / Gmail\napproves", GREEN),
        ("4. Dispatched", "Shipping\ninvoice + dispatch", BLUE),
    ]
    sw, sh = 170, 110
    gap = 36
    total = len(steps) * sw + (len(steps) - 1) * gap
    x = (w - total) // 2
    y = 70

    for i, (title, body, color) in enumerate(steps):
        draw_rounded_rect(draw, (x, y, x + sw, y + sh), fill=color, radius=12)
        draw.text((x + sw // 2, y + 22), title, fill=WHITE if color == BLUE else "#222", font=box_f, anchor="mm")
        draw.multiline_text((x + sw // 2, y + 62), body, fill="#222" if color != BLUE else WHITE, font=try_font(12), anchor="mm", align="center")
        if i < len(steps) - 1:
            ax = x + sw + 6
            draw.polygon([(ax, y + sh // 2 - 10), (ax + gap - 12, y + sh // 2), (ax, y + sh // 2 + 10)], fill=BLUE)
        x += sw + gap

    draw.text((w // 2, 210), "Side paths: Rejected · Canceled · Email Failed (admin retries or approves manually)", fill="#666", font=try_font(12), anchor="mm")
    return save_diagram("order-flow.png", img)


def diagram_roles() -> Path:
    w, h = 900, 340
    img = Image.new("RGB", (w, h), WHITE)
    draw = ImageDraw.Draw(img)
    title_f = try_font(22, True)
    box_f = try_font(14, True)
    small_f = try_font(12)

    draw.text((w // 2, 24), "Who does what", fill=BLUE, font=title_f, anchor="mm")

    roles = [
        ("Admin", "/admin", "Targets · Rates · Orders\nApprove · Reports · Team", BLUE, WHITE),
        ("Shipping", "/shipping", "Approved queue\nInvoice · Transport · Dispatch", "#00838f", WHITE),
        ("Distributor", "/distributor", "Place orders · View targets\nPhysical stock · Prices", RED, WHITE),
        ("Platform", "/platform", "All workspaces\nGmail API credentials", "#5e35b1", WHITE),
    ]
    bw, bh = 190, 200
    gap = 24
    start = (w - (4 * bw + 3 * gap)) // 2
    y = 60
    for i, (name, path, body, bg, fg) in enumerate(roles):
        x = start + i * (bw + gap)
        draw_rounded_rect(draw, (x, y, x + bw, y + bh), fill=bg, radius=14)
        draw.text((x + bw // 2, y + 30), name, fill=fg, font=box_f, anchor="mm")
        draw.text((x + bw // 2, y + 58), path, fill=fg, font=small_f, anchor="mm")
        draw.multiline_text((x + bw // 2, y + 120), body, fill=fg, font=small_f, anchor="mm", align="center")

    return save_diagram("roles.png", img)


def diagram_setup() -> Path:
    w, h = 900, 300
    img = Image.new("RGB", (w, h), WHITE)
    draw = ImageDraw.Draw(img)
    title_f = try_font(22, True)
    phase_f = try_font(14, True)
    body_f = try_font(12)

    draw.text((w // 2, 24), "Setup timeline (first-time)", fill=BLUE, font=title_f, anchor="mm")

    phases = [
        ("Phase 1\nPlatform", "SQL migrations\nGmail Client ID + API Key\nDeploy app", BLUE),
        ("Phase 2\nSign up", "Create workspace\nOwner account\nOnboarding wizard", RED),
        ("Phase 3\nConfigure", "Distributors · Rates\nTargets · Schemes\nTeam invites", GREEN),
        ("Phase 4\nGo live", "Connect Gmail\nDistributors sign in\nOrders → dispatch", AMBER),
    ]
    pw, ph = 180, 190
    gap = 28
    start = (w - (4 * pw + 3 * gap)) // 2
    y = 60
    for i, (title, body, color) in enumerate(phases):
        x = start + i * (pw + gap)
        draw_rounded_rect(draw, (x, y, x + pw, y + ph), fill=color, radius=12)
        draw.multiline_text((x + pw // 2, y + 36), title, fill=WHITE if color in (BLUE, RED) else "#222", font=phase_f, anchor="mm", align="center")
        draw.multiline_text((x + pw // 2, y + 115), body, fill=WHITE if color in (BLUE, RED) else "#222", font=body_f, anchor="mm", align="center")
        if i < 3:
            ax = x + pw + 4
            draw.polygon([(ax, y + ph // 2 - 8), (ax + gap - 8, y + ph // 2), (ax, y + ph // 2 + 8)], fill=BLUE)

    return save_diagram("setup-timeline.png", img)


def diagram_admin_modules() -> Path:
    w, h = 900, 420
    img = Image.new("RGB", (w, h), WHITE)
    draw = ImageDraw.Draw(img)
    title_f = try_font(22, True)
    mod_f = try_font(12, True)

    draw.text((w // 2, 24), "Admin dashboard — main modules", fill=BLUE, font=title_f, anchor="mm")

    modules = [
        "Dashboard", "Orders", "Calculator", "Targets", "Scheme & Discount",
        "Product & Rate Master", "Inventory", "Physical Stock", "Stock lifting",
        "Distributors", "Reports", "Activity", "GST Settings", "Workspace",
        "Team & invites", "User & Permissions",
    ]
    cols, rows = 4, 4
    mw, mh, gap = 200, 72, 16
    start_x = (w - (cols * mw + (cols - 1) * gap)) // 2
    start_y = 60
    colors = ["#e3f2fd", "#fce4ec", "#e8f5e9", "#fff3e0"]
    for idx, name in enumerate(modules):
        r, c = divmod(idx, cols)
        x = start_x + c * (mw + gap)
        y = start_y + r * (mh + gap)
        draw_rounded_rect(draw, (x, y, x + mw, y + mh), fill=colors[c % 4], outline=BLUE_LIGHT)
        draw.text((x + mw // 2, y + mh // 2), name, fill="#222", font=mod_f, anchor="mm")

    return save_diagram("admin-modules.png", img)


def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_document_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)


def add_centered_image(doc: Document, path: Path, width_in: float = 1.4) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))


def add_figure(doc: Document, path: Path, caption: str, width_in: float = 6.2) -> None:
    add_centered_image(doc, path, width_in)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(10)
    cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()


def add_title_page(doc: Document, logo_path: Path) -> None:
    doc.add_paragraph()
    add_centered_image(doc, logo_path, 1.6)
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Distribution Prime")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Standard Operating Procedure\nVisual Summary Guide")
    r2.font.size = Pt(16)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version 1.2  |  Effective {date.today().strftime('%B %d, %Y')}\n").bold = True
    meta.add_run("Application: https://distribution-prime.pages.dev\n")
    meta.add_run(
        "Professional distribution management for bottlers and distributor networks\n"
        "Isolated workspaces · Orders · Targets · Shipping · Reports"
    )

    doc.add_page_break()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "0D47A1") -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], header_fill)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def add_flow_status_table(doc: Document) -> None:
    headers = ["Step", "Status", "Who", "Action"]
    rows = [
        ["1", "Pending", "Distributor → Admin", "Distributor submits order; appears in Admin Orders queue"],
        ["2", "Sent / Email Failed", "Admin", "Send approval email via Gmail; retry if failed"],
        ["3", "Approved / Rejected", "Admin / GM", "Approve manually or via Gmail reply keywords"],
        ["4", "Dispatched", "Shipping", "Upload invoice, transport details, click Dispatch"],
        ["—", "Canceled", "Distributor", "Cancel while Pending or Sent only"],
    ]
    add_table(doc, headers, rows)


def build_document(visuals: dict[str, Path], logo_path: Path) -> Document:
    doc = Document()
    set_document_defaults(doc)
    add_title_page(doc, logo_path)

    # --- 1 Purpose ---
    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This Visual Summary SOP explains how to operate Distribution Prime end-to-end: creating a company "
        "workspace, configuring master data, running the order approval and dispatch workflow, and using "
        "reports and stock tools. It is intended for platform operators, workspace admins, shipping staff, "
        "and distributors."
    )
    doc.add_paragraph(
        "Each organization (workspace) has its own isolated data. Users always sign in with a Workspace ID "
        "so they only see their company's distributors, orders, targets, and reports."
    )

    # --- 2 Scope ---
    doc.add_heading("2. Scope", level=1)
    add_table(
        doc,
        ["In scope", "Out of scope"],
        [
            ["Workspace sign-up and branding", "Supabase / server administration"],
            ["Admin, Shipping, Distributor dashboards", "Custom software development"],
            ["Orders, targets, rates, schemes, GST", "Unless you are a platform operator: /platform console internals"],
            ["Inventory, physical stock, stock lifting", ""],
            ["Reports, activity log, team invites", ""],
            ["Gmail for order approval emails", ""],
        ],
        header_fill="1565C0",
    )

    # --- 3 Architecture visual ---
    doc.add_heading("3. System overview — multi-tenant architecture", level=1)
    doc.add_paragraph(
        "Distribution Prime uses one shared database. Every table row (orders, distributors, targets, etc.) "
        "includes an organization_id. Supabase Row Level Security (RLS) plus application filters prevent "
        "cross-company data access."
    )
    add_figure(doc, visuals["architecture"], "Figure 1 — Each workspace is logically separate inside one platform.")

    add_bullets(
        doc,
        [
            "Workspace ID: short login name (e.g. acme-beverages). Login URL: /w/{workspace-id}/login",
            "Admin & shipping sign in with email + password. Distributors sign in with code + password.",
            "Platform operator manages all workspaces at /platform and sets Gmail API credentials for all orgs.",
        ],
    )

    # --- 4 Roles visual ---
    doc.add_heading("4. User roles and dashboards", level=1)
    add_figure(doc, visuals["roles"], "Figure 2 — Role-based dashboards and responsibilities.")
    add_table(
        doc,
        ["Role", "Sign-in credentials", "Dashboard path", "Key tasks"],
        [
            ["Admin", "Workspace ID + email + password", "/admin", "Configure workspace, approve orders, manage rates/targets/team"],
            ["Viewer", "Same as admin", "/admin", "Read-only access to dashboards and reports"],
            ["Shipping", "Workspace ID + email + password", "/shipping", "Invoice upload, transport entry, dispatch"],
            ["Distributor", "Workspace ID + code + password", "/distributor", "Place orders, view targets/prices, physical stock"],
            ["Platform admin", "Email at /platform/login", "/platform", "Manage workspaces; configure Gmail API for all tenants"],
        ],
    )

    # --- 5 Setup visual ---
    doc.add_heading("5. Initial setup — visual timeline", level=1)
    add_figure(doc, visuals["setup"], "Figure 3 — Recommended setup sequence from platform to go-live.")
    doc.add_heading("5.1 Platform operator (once per environment)", level=2)
    add_numbered(
        doc,
        [
            "Deploy app with REACT_APP_SUPABASE_URL and REACT_APP_SUPABASE_ANON_KEY.",
            "Run SQL migrations: tenant RLS, workspace signup RPC, platform admin, platform Gmail credentials.",
            "Sign in at /platform/login.",
            "Platform console → Gmail → save Google OAuth Client ID and API Key (copied to every workspace).",
        ],
    )
    doc.add_heading("5.2 Organization owner — create workspace", level=2)
    add_numbered(
        doc,
        [
            "Open /signup → enter company name, workspace ID, owner email, password (min. 8 chars).",
            "Optional: invoice letterhead (address, post no., GST) in collapsible section.",
            "Complete onboarding wizard → Authorize Gmail (connect once per browser/device).",
            "Share branded login link: https://distribution-prime.pages.dev/w/{workspace-id}/login",
        ],
    )
    doc.add_heading("5.3 Admin — configure workspace", level=2)
    add_figure(doc, visuals["admin_modules"], "Figure 4 — Admin sidebar modules (configuration and operations).", width_in=6.0)
    add_numbered(
        doc,
        [
            "Workspace — company name, address, GST, theme color, login URL.",
            "Distributors — add codes, regions, passwords; import bulk if needed.",
            "Product & Rate Master — SKUs, CSD/Water/CAN categories, rates, UC divisor.",
            "Targets — set global target period; enter PC and UC targets per distributor.",
            "Scheme & Discount — free cases or % discount by SKU/category and date range.",
            "Team & invites — invite admin, viewer, shipping; User & Permissions for roles.",
        ],
    )

    # --- 6 Order flow visual ---
    doc.add_heading("6. Order lifecycle", level=1)
    add_figure(doc, visuals["order_flow"], "Figure 5 — Standard order path from placement to dispatch.")
    add_flow_status_table(doc)
    doc.add_paragraph(
        "On Dispatch: distributor target achievement (PC/UC) updates automatically and stock lifting "
        "records are created. Distributor can download the shipping invoice from their Orders view."
    )

    # --- 7 Daily ops ---
    doc.add_heading("7. Daily operating procedures", level=1)

    doc.add_heading("7.1 Admin — morning checklist", level=2)
    add_table(
        doc,
        ["Step", "Module", "Action"],
        [
            ["1", "Dashboard", "Review regional/distributor performance vs targets"],
            ["2", "Orders", "Process Pending and Sent tabs — view, email, approve, reject"],
            ["3", "Physical Stock", "Review new distributor stock submissions (badge indicator)"],
            ["4", "Stock lifting / Reports", "Export or review dispatched sales data"],
            ["5", "Activity", "Check audit log for errors or unusual actions"],
            ["6", "Gmail chip (app bar)", "Ensure Gmail shows connected before sending order emails"],
        ],
        header_fill="2E7D32",
    )

    doc.add_heading("7.2 Distributor — order day", level=2)
    add_table(
        doc,
        ["Step", "Screen", "Action"],
        [
            ["1", "Home", "Check target balance (CSD/Water PC & UC) and days remaining"],
            ["2", "Place Order", "Select products, enter cases; schemes apply automatically"],
            ["3", "Orders", "Confirm new order shows as Pending"],
            ["4", "Stock", "Submit physical stock by SKU and MFG date if required"],
            ["5", "Prices", "Reference read-only rate list from Rate Master"],
        ],
        header_fill="E40521",
    )

    doc.add_heading("7.3 Shipping — dispatch day", level=2)
    add_table(
        doc,
        ["Step", "Requirement", "Details"],
        [
            ["1", "Order status", "Must be Approved"],
            ["2", "Invoice", "Upload shipping invoice (PDF/image)"],
            ["3", "Transport", "Transporter, vehicle type, vehicle number, charges"],
            ["4", "Dispatch", "Click Dispatch → status Dispatched; achievement updates"],
        ],
        header_fill="00838F",
    )

    # --- 8 Gmail ---
    doc.add_heading("8. Gmail integration (optional)", level=1)
    add_table(
        doc,
        ["Layer", "Who configures", "What it does"],
        [
            ["API credentials", "Platform operator", "Client ID + API Key saved in Platform console → applies to all workspaces"],
            ["OAuth connection", "Each workspace admin", "Authorize Gmail once per browser — sends order emails from admin's Gmail"],
            ["Status indicator", "Admin app bar", "Shows Gmail connected / not configured / connect prompt"],
            ["Auto-approval", "System (optional)", "Reads Gmail replies for approve/reject keywords after Send Email"],
        ],
    )
    add_bullets(
        doc,
        [
            "Google Cloud: add app URL to Authorized JavaScript origins.",
            "Testing mode: add admin emails as Test users in OAuth consent screen.",
            "Allow browser popups when clicking Authorize Gmail.",
        ],
    )

    # --- 9 Glossary ---
    doc.add_heading("9. Key terms", level=1)
    add_table(
        doc,
        ["Term", "Definition"],
        [
            ["PC (Physical Case)", "Case count by product category; CAN products excluded from PC totals"],
            ["UC (Unit Case)", "Normalized volume: (cases × UC multiplier) ÷ UC divisor — used for targets"],
            ["CSD / Water / CAN", "Product categories in Rate Master"],
            ["Stock lifting", "Sales volume recorded automatically when shipping dispatches an order"],
            ["Physical stock", "Distributor-reported on-hand inventory by SKU and MFG date"],
            ["Workspace", "Isolated company tenant; all data scoped by organization_id"],
            ["Dispatched", "Final shipment status in UI (stored as delivered in database)"],
        ],
    )

    # --- 10 Troubleshooting ---
    doc.add_heading("10. Troubleshooting", level=1)
    add_table(
        doc,
        ["Symptom", "Likely cause", "Fix"],
        [
            ["Signup fails — email exists", "Email already in Supabase Auth", "Sign in instead or use new email"],
            ["No access to workspace", "Wrong workspace ID or no admin row", "Use correct workspace ID; re-signup or invite"],
            ["Gmail not configured", "Platform Gmail not saved", "Platform console → Gmail → save credentials"],
            ["Gmail popup timeout", "Popups blocked", "Allow popups for site; retry Authorize Gmail"],
            ["Google unverified app", "OAuth app not verified / not test user", "Add test users or submit Google verification"],
            ["Cannot dispatch", "Missing invoice or transport", "Complete all shipping fields on Approved order"],
            ["Distributor login fails", "Wrong code or password", "Admin resets in Distributors dialog"],
        ],
        header_fill="5E35B1",
    )

    # --- 11 Security ---
    doc.add_heading("11. Security and data handling", level=1)
    add_bullets(
        doc,
        [
            "Never share passwords between users. Distributors use codes; staff use email accounts.",
            "Sign out on shared computers (admin/shipping). Distributor sessions may persist if Remember me is used.",
            "Workspace data isolation enforced by organization_id + Supabase RLS.",
            "Gmail OAuth tokens stored in browser localStorage on connected devices only.",
            "Platform operators can access all workspaces for support — restrict platform admin accounts.",
        ],
    )

    # --- 12 Document control ---
    doc.add_heading("12. Document control", level=1)
    add_table(
        doc,
        ["Version", "Date", "Changes"],
        [
            ["1.0", "March 2025", "Initial SOP"],
            ["1.1", "June 2025", "Summary SOP with platform Gmail and multi-tenant notes"],
            ["1.2", date.today().strftime("%B %Y"), "Visual summary with logo, diagrams, expanded role/setup/order sections"],
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_centered_image(doc, logo_path, 0.9)
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Distribution Prime — Visual Summary SOP\n").bold = True
    footer.add_run("Support: codewynbuild@gmail.com\n")
    footer.add_run("Privacy: https://distribution-prime.pages.dev/legal/privacy-policy")

    return doc


def main() -> None:
    logo_path = resolve_logo()
    visuals = {
        "architecture": diagram_architecture(),
        "order_flow": diagram_order_flow(),
        "roles": diagram_roles(),
        "setup": diagram_setup(),
        "admin_modules": diagram_admin_modules(),
    }
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document(visuals, logo_path)
    doc.save(OUT)
    print(f"Wrote {OUT}")
    print("Visual assets:", ", ".join(str(v) for v in visuals.values()))


if __name__ == "__main__":
    main()
